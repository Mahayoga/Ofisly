from docx import Document

document = Document('template.docx')

# for table in document.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             for para in cell.paragraphs:
#                 if '__' in para.text:
#                     for run in para.runs:
#                         run.bold = True
#                         print(para.text)

table = document.tables[1]
cell1 = table.rows[3].cells[0]
cell2 = table.rows[4].cells[0]

for para in cell1.paragraphs:
    for run in para.runs:
        run.bold = True
        run.underline = True
        print(para.text)

for para in cell2.paragraphs:
    for run in para.runs:
        run.italic = True
        print(para.text)    
# document.save('template1.docx')
