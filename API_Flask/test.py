# import pythoncom
# import comtypes.client
# import os
# import threading
# import requests
# import re
# import traceback
# import json
# from flask import Flask, jsonify, request
# from flask_cors import CORS
# from dotenv import load_dotenv
# from datetime import datetime
# from docx import Document

# # Inisialisasi
# app = Flask(__name__)
# CORS(app)
# load_dotenv()
# wdFormatPDF = 17
# bulan = ['Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']
# processesId = [] # Daftar untuk melacak ID yang sedang diproses

# # ======================================================================
# # FUNGSI BACKGROUND
# # Tugas: Membuat file, mengirimnya ke Laravel, lalu membersihkan.
# # ======================================================================
# def background_task(filename_docx, filename_pdf, laravel_upload_url, surat_id):
#     """Fungsi ini berjalan di background untuk mengunggah file dan membersihkan."""
#     try:
#         # 1. Mengirim file ke endpoint khusus di Laravel
#         print(f"[{surat_id}] Mengirim file DOCX & PDF ke Laravel di: {laravel_upload_url}")
#         with open(filename_docx, 'rb') as f_docx, open(filename_pdf, 'rb') as f_pdf:
#             files = {
#                 'file_docx': (os.path.basename(filename_docx), f_docx, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
#                 'file_pdf': (os.path.basename(filename_pdf), f_pdf, 'application/pdf')
#             }
#             payload = {'surat_id': surat_id}
            
#             res = requests.post(laravel_upload_url, files=files, data=payload, timeout=30)
#             res.raise_for_status() # Akan error jika status code bukan 2xx
            
#             response_json = res.json()
#             if response_json.get('success'):
#                 print(f"[{surat_id}] Upload ke Laravel berhasil.")
#             else:
#                 print(f"[{surat_id}] Laravel merespon dengan error: {response_json.get('message')}")

#     except requests.exceptions.RequestException as e:
#         print(f"[{surat_id}] KRITIS: Gagal mengirim file ke Laravel: {e}")
#     except Exception as e:
#         print(f"[{surat_id}] KRITIS: Terjadi error di background thread: {e}")
#     finally:
#         # 2. Selalu membersihkan file sementara dan status proses
#         try:
#             if os.path.exists(filename_docx): os.remove(filename_docx)
#             if os.path.exists(filename_pdf): os.remove(filename_pdf)
#             print(f"[{surat_id}] File sementara di Flask telah dihapus.")
#             if surat_id in processesId:
#                 processesId.remove(surat_id)
#         except (ValueError, OSError) as e:
#             print(f"[{surat_id}] Gagal membersihkan file/proses: {e}")
        
#         print(f"[{surat_id}] Proses di sisi Flask selesai.")


# # ======================================================================
# # ENDPOINT UTAMA: GENERATE SURAT (SUDAH DIPERBAIKI)
# # ======================================================================
# @app.route('/generate/surat/promotor', methods=['POST'])
# def generate_surat_promotor():
#     # URL Laravel untuk MENERIMA file yang sudah jadi
#     laravel_upload_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/surat-promotor/upload-final'
#     file_template_path = os.getenv('TEMPLATE_PROMOTOR_PATH')
    
#     word = doc = None
#     id_surat = "UNKNOWN"

#     try:
#         # 1. Validasi Input
#         if not request.json or 'surat_data' not in request.json:
#             return jsonify({'status': 'error', 'message': 'Payload JSON atau surat_data tidak ditemukan.'}), 400
        
#         surat_data = request.json['surat_data']
#         id_surat = str(surat_data.get('id_surat_tugas_promotor', 'UNKNOWN'))

#         if id_surat in processesId:
#             print(f"[{id_surat}] Peringatan: Proses untuk ID ini sudah berjalan.")
#             return jsonify({'status': 'processing', 'message': 'Proses untuk ID ini sudah berjalan.'})

#         print(f"[{id_surat}] Memulai proses pembuatan surat...")
#         pythoncom.CoInitialize()

#         # 2. Proses Data (Fungsi dan Logika Anda sudah benar)
#         def format_date(date_str):
#             """Fungsi untuk memformat tanggal dari ISO format ke format Indonesia."""
#             if not date_str: 
#                 return ""
#             try:
#                 if isinstance(date_str, str):
#                     if 'T' in date_str:
#                         date_str = date_str.split('T')[0]  # Ambil bagian date saja (YYYY-MM-DD)
                    
#                     date_obj = datetime.strptime(date_str, '%Y-%m-%d')
#                 else:
#                     date_obj = date_str
                    
#                 return f"{date_obj.day} {bulan[date_obj.month - 1]} {date_obj.year}"
#             except (ValueError, TypeError, IndexError) as e:
#                 print(f"Error formatting date: {date_str}, Error: {e}")
#                 return str(date_str) # Kembalikan string asli jika gagal format

#         penempatan_raw = surat_data.get('penempatan', '[]')
#         penempatan_list = []
#         try:
#             # Dengan $casts di Laravel, ini akan selalu berupa array/list
#             parsed_data = json.loads(penempatan_raw) if isinstance(penempatan_raw, str) else penempatan_raw
#             if isinstance(parsed_data, list):
#                  penempatan_list = [str(item).strip() for item in parsed_data if str(item).strip()]
#             else:
#                  penempatan_list = [str(parsed_data)]
#         except (json.JSONDecodeError, TypeError):
#             penempatan_list = [item.strip() for item in str(penempatan_raw).split(',') if item.strip()]

#         penempatan_text = "-"
#         if penempatan_list:
#             penempatan_text = penempatan_list[0] if len(penempatan_list) == 1 else "\n".join(f"{i}. {item}" for i, item in enumerate(penempatan_list, 1))

#         nama_kandidat = surat_data.get('nama_kandidat', '-').strip() or '-'
#         print(f"[{id_surat}] Data diterima untuk kandidat: {nama_kandidat}")

#         # 3. Manipulasi Dokumen Word
#         if not os.path.exists(file_template_path):
#             raise FileNotFoundError(f"Template tidak ditemukan di: {file_template_path}")
#         document = Document(file_template_path)
        
#         # ============================================================
#         # PERBAIKAN UTAMA DI SINI: Panggil format_date() untuk setiap tanggal
#         # ============================================================
#         replacements = {
#             '__PENEMPATAN__': penempatan_text,
#             '__DATATGLSURATPEMBUATAN__': format_date(surat_data.get('tgl_surat_pembuatan')),
#             '__NAMAKANDIDAT__': nama_kandidat,
#             '__DATATGLPENUGASAN__': format_date(surat_data.get('tgl_penugasan')), # <-- INI PERBAIKANNYA
#         }
#         # ============================================================
        
#         for key, value in replacements.items():
#             for p in document.paragraphs: p.text = p.text.replace(key, str(value))
#             for t in document.tables:
#                 for r in t.rows:
#                     for c in r.cells: c.text = c.text.replace(key, str(value))
        
#         # ... (sisa kode Anda untuk menyimpan file dan memulai thread sudah benar) ...
#         # 4. Simpan & Konversi ke PDF
#         base_filename = f'Surat_Penempatan_Promotor_{nama_kandidat}_{id_surat}'
#         safe_filename = re.sub(r'[^\w\-_\. ]', '_', base_filename)
#         dirname = os.path.dirname(os.path.abspath(__file__))
#         temp_dir = os.path.join(dirname, 'temp')
#         os.makedirs(temp_dir, exist_ok=True) # Buat folder temp jika belum ada
#         filename_docx = os.path.join(temp_dir, f'{safe_filename}.docx')
#         filename_pdf = os.path.join(temp_dir, f'{safe_filename}.pdf')
#         document.save(filename_docx)

#         print(f"[{id_surat}] Mengonversi ke PDF...")
#         word = comtypes.client.CreateObject('Word.Application')
#         word.Visible = False
#         doc = word.Documents.Open(os.path.abspath(filename_docx), ReadOnly=True)
#         doc.SaveAs(os.path.abspath(filename_pdf), FileFormat=wdFormatPDF)
#         doc.Close(); doc = None
#         word.Quit(); word = None
#         print(f"[{id_surat}] Konversi PDF berhasil.")

#         # 5. Jalankan Proses Background untuk Upload & Cleanup
#         processesId.append(id_surat)
#         upload_thread = threading.Thread(
#             target=background_task,
#             args=(filename_docx, filename_pdf, laravel_upload_url, id_surat)
#         )
#         upload_thread.daemon = True
#         upload_thread.start()
        
#         print(f"[{id_surat}] Proses pembuatan file di Flask berhasil dimulai dan diserahkan ke background thread.")
#         return jsonify({'status': 'success', 'message': 'Proses pembuatan file dimulai.'})

#     except Exception as e:
#         print(f"[{id_surat}] ERROR FATAL: {e}")
#         traceback.print_exc()
#         if id_surat in processesId: processesId.remove(id_surat)
#         return jsonify({'status': 'error', 'message': f'Terjadi kesalahan internal di Flask: {e}'}), 500

#     finally:
#         # 6. Cleanup COM
#         try:
#             if doc: doc.Close()
#             if word: word.Quit()
#             pythoncom.CoUninitialize()
#             print(f"[{id_surat}] Cleanup COM berhasil.")
#         except Exception as cleanup_error:
#             print(f"[{id_surat}] Error saat cleanup COM: {cleanup_error}")

# # ======================================================================
# # ENDPOINT PENGECEKAN
# # Memberi tahu Laravel apakah sebuah proses masih berjalan di Flask.
# # ======================================================================
# @app.route('/check/generate/run', methods=['POST'])
# def check_generate_files():
#     req_id = request.json.get('id')
#     if req_id and str(req_id) in processesId:
#         return jsonify({'status': True})
#     return jsonify({'status': False})

# if __name__ == '__main__':
#     app.run(debug=True, port=5000)

# import pythoncom
# import comtypes.client
# import os
# import threading
# import requests
# import re
# import traceback
# import json
# from flask import Flask, jsonify, request
# from flask_socketio import SocketIO, emit
# from flask_cors import CORS
# from dotenv import load_dotenv
# from datetime import datetime
# from docx import Document

# # Inisialisasi
# app = Flask(__name__)
# CORS(app, origins=["http://localhost:8000", "http://127.0.0.1:8000"])

# # Inisialisasi SocketIO
# socketio = SocketIO(app, 
#                    cors_allowed_origins=["http://localhost:8000", "http://127.0.0.1:8000"])

# load_dotenv()
# wdFormatPDF = 17
# bulan = ['Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']

# # Dictionary untuk melacak status proses
# process_status = {}

# # ======================================================================
# # FUNGSI WEB SOCKET EVENT HANDLERS
# # ======================================================================
# @socketio.on('connect')
# def handle_connect():
#     print('Client connected:', request.sid)
#     emit('connected', {'message': 'Connected to Flask WebSocket'})

# @socketio.on('disconnect')
# def handle_disconnect():
#     print('Client disconnected:', request.sid)

# @socketio.on('subscribe_to_progress')
# def handle_subscribe(data):
#     surat_id = data.get('surat_id')
#     if surat_id:
#         socketio.enter_room(request.sid, f'surat_{surat_id}')
#         emit('subscribed', {'surat_id': surat_id}, room=request.sid)
#         print(f'Client {request.sid} subscribed to surat_id: {surat_id}')

# # ======================================================================
# # FUNGSI UNTUK MENGIRIM STATUS MELALUI WEB SOCKET
# # ======================================================================
# def send_status(surat_id, status, message=None, progress=None, file_paths=None):
#     """Mengirim status update ke client melalui WebSocket"""
#     data = {
#         'surat_id': surat_id,
#         'status': status,
#         'timestamp': datetime.now().isoformat()
#     }
    
#     if message:
#         data['message'] = message
#     if progress is not None:
#         data['progress'] = progress
#     if file_paths:
#         data['file_paths'] = file_paths
    
#     # Update status terbaru
#     process_status[surat_id] = data
    
#     # Kirim ke room yang sesuai
#     socketio.emit('generation_status', data, room=f'surat_{surat_id}')
#     print(f"[{surat_id}] Status update: {status} - {message}")

# # ======================================================================
# # FUNGSI BACKGROUND TASK DENGAN PROGRESS UPDATE
# # ======================================================================
# def background_task(surat_data):
#     """Fungsi background dengan progress reporting melalui WebSocket"""
#     surat_id = str(surat_data.get('id_surat_tugas_promotor', 'UNKNOWN'))
#     laravel_upload_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/surat-promotor/upload-final'
    
#     filename_docx = None
#     filename_pdf = None
    
#     try:
#         # Update status: Memulai proses
#         send_status(surat_id, 'processing', 'Memulai proses pembuatan surat...', 10)
        
#         # 1. Persiapan data
#         pythoncom.CoInitialize()
        
#         def format_date(date_str):
#             """Fungsi untuk memformat tanggal dari ISO format ke format Indonesia."""
#             if not date_str: 
#                 return ""
#             try:
#                 if isinstance(date_str, str):
#                     if 'T' in date_str:
#                         date_str = date_str.split('T')[0]
#                     date_obj = datetime.strptime(date_str, '%Y-%m-%d')
#                 else:
#                     date_obj = date_str
#                 return f"{date_obj.day} {bulan[date_obj.month - 1]} {date_obj.year}"
#             except (ValueError, TypeError, IndexError) as e:
#                 print(f"Error formatting date: {date_str}, Error: {e}")
#                 return str(date_str)

#         # Process penempatan
#         penempatan_raw = surat_data.get('penempatan', '[]')
#         penempatan_list = []
#         try:
#             parsed_data = json.loads(penempatan_raw) if isinstance(penempatan_raw, str) else penempatan_raw
#             if isinstance(parsed_data, list):
#                  penempatan_list = [str(item).strip() for item in parsed_data if str(item).strip()]
#             else:
#                  penempatan_list = [str(parsed_data)]
#         except (json.JSONDecodeError, TypeError):
#             penempatan_list = [item.strip() for item in str(penempatan_raw).split(',') if item.strip()]

#         penempatan_text = "-"
#         if penempatan_list:
#             penempatan_text = penempatan_list[0] if len(penempatan_list) == 1 else "\n".join(f"{i}. {item}" for i, item in enumerate(penempatan_list, 1))

#         nama_kandidat = surat_data.get('nama_kandidat', '-').strip() or '-'
        
#         send_status(surat_id, 'processing', 'Memproses data template...', 30)
        
#         # 2. Manipulasi dokumen Word
#         file_template_path = os.getenv('TEMPLATE_PROMOTOR_PATH')
#         if not os.path.exists(file_template_path):
#             raise FileNotFoundError(f"Template tidak ditemukan di: {file_template_path}")
        
#         document = Document(file_template_path)
        
#         replacements = {
#             '__PENEMPATAN__': penempatan_text,
#             '__DATATGLSURATPEMBUATAN__': format_date(surat_data.get('tgl_surat_pembuatan')),
#             '__NAMAKANDIDAT__': nama_kandidat,
#             '__DATATGLPENUGASAN__': format_date(surat_data.get('tgl_penugasan')),
#         }
        
#         for key, value in replacements.items():
#             for p in document.paragraphs: 
#                 p.text = p.text.replace(key, str(value))
#             for t in document.tables:
#                 for r in t.rows:
#                     for c in r.cells: 
#                         c.text = c.text.replace(key, str(value))
        
#         send_status(surat_id, 'processing', 'Menyimpan dokumen Word...', 50)
        
#         # 3. Simpan file DOCX
#         base_filename = f'Surat_Penempatan_Promotor_{nama_kandidat}_{surat_id}'
#         safe_filename = re.sub(r'[^\w\-_\. ]', '_', base_filename)
#         dirname = os.path.dirname(os.path.abspath(__file__))
#         temp_dir = os.path.join(dirname, 'temp')
#         os.makedirs(temp_dir, exist_ok=True)
#         filename_docx = os.path.join(temp_dir, f'{safe_filename}.docx')
#         filename_pdf = os.path.join(temp_dir, f'{safe_filename}.pdf')
        
#         document.save(filename_docx)
        
#         send_status(surat_id, 'processing', 'Mengonversi ke PDF...', 70)
        
#         # 4. Konversi ke PDF
#         word = comtypes.client.CreateObject('Word.Application')
#         word.Visible = False
#         doc = word.Documents.Open(os.path.abspath(filename_docx), ReadOnly=True)
#         doc.SaveAs(os.path.abspath(filename_pdf), FileFormat=wdFormatPDF)
#         doc.Close()
#         word.Quit()
        
#         send_status(surat_id, 'processing', 'Mengupload file ke server...', 90)
        
#         # 5. Upload ke Laravel
#         with open(filename_docx, 'rb') as f_docx, open(filename_pdf, 'rb') as f_pdf:
#             files = {
#                 'file_docx': (os.path.basename(filename_docx), f_docx, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
#                 'file_pdf': (os.path.basename(filename_pdf), f_pdf, 'application/pdf')
#             }
#             payload = {'surat_id': surat_id}
            
#             res = requests.post(laravel_upload_url, files=files, data=payload, timeout=30)
#             res.raise_for_status()
            
#             response_json = res.json()
#             if response_json.get('success'):
#                 send_status(surat_id, 'completed', 'File berhasil digenerate dan diupload!', 100)
#             else:
#                 raise Exception(f"Laravel error: {response_json.get('message')}")

#     except Exception as e:
#         error_msg = f"Error dalam proses: {str(e)}"
#         print(f"[{surat_id}] ERROR: {error_msg}")
#         traceback.print_exc()
#         send_status(surat_id, 'error', error_msg)
        
#     finally:
#         # 6. Cleanup
#         try:
#             if filename_docx and os.path.exists(filename_docx): 
#                 os.remove(filename_docx)
#             if filename_pdf and os.path.exists(filename_pdf): 
#                 os.remove(filename_pdf)
#             pythoncom.CoUninitialize()
#         except Exception as cleanup_error:
#             print(f"[{surat_id}] Cleanup error: {cleanup_error}")

# # ======================================================================
# # ENDPOINT UTAMA DENGAN WEB SOCKET
# # ======================================================================
# @app.route('/generate/surat/promotor', methods=['POST'])
# def generate_surat_promotor():
#     try:
#         if not request.json or 'surat_data' not in request.json:
#             return jsonify({'status': 'error', 'message': 'Payload JSON atau surat_data tidak ditemukan.'}), 400
        
#         surat_data = request.json['surat_data']
#         surat_id = str(surat_data.get('id_surat_tugas_promotor', 'UNKNOWN'))
        
#         # Cek apakah proses sudah berjalan untuk surat_id ini
#         if surat_id in process_status and process_status[surat_id].get('status') == 'processing':
#             return jsonify({
#                 'status': 'processing', 
#                 'message': 'Proses untuk ID ini sudah berjalan.',
#                 'websocket_required': True
#             })
        
#         # Mulai background task tanpa parameter sid
#         socketio.start_background_task(background_task, surat_data)
        
#         return jsonify({
#             'status': 'success', 
#             'message': 'Proses pembuatan file dimulai.',
#             'surat_id': surat_id,
#             'websocket_required': True
#         })
        
#     except Exception as e:
#         error_msg = f"Error memulai proses: {str(e)}"
#         print(f"ERROR in generate_surat_promotor: {error_msg}")
#         return jsonify({'status': 'error', 'message': error_msg}), 500

# # ======================================================================
# # ENDPOINT UNTUK MENGECEK STATUS (Fallback)
# # ======================================================================
# @app.route('/check/status/<surat_id>', methods=['GET'])
# def check_status(surat_id):
#     """Endpoint fallback untuk cek status (jika WebSocket tidak available)"""
#     status = process_status.get(surat_id, {'status': 'not_found'})
#     return jsonify(status)

# # ======================================================================
# # ENDPOINT UNTUK MENDAPATKAN HISTORY STATUS
# # ======================================================================
# @app.route('/status/history', methods=['GET'])
# def status_history():
#     """Mendapatkan history status semua proses"""
#     return jsonify(process_status)

# if __name__ == '__main__':
#     print("Starting Flask server with WebSocket support...")
#     socketio.run(app, debug=True, port=5000, host='0.0.0.0')

import os
import threading
import requests
import re
import traceback
import json
from flask import Flask, jsonify, request
from flask_socketio import SocketIO, emit
from flask_cors import CORS
from dotenv import load_dotenv
from datetime import datetime
from docx import Document
from docx2pdf import convert
import pythoncom

# Inisialisasi
app = Flask(__name__)
CORS(app, origins=["http://localhost:8000", "http://127.0.0.1:8000"])

# Inisialisasi SocketIO
socketio = SocketIO(app, 
                   cors_allowed_origins=["http://localhost:8000", "http://127.0.0.1:8000"])

load_dotenv()
bulan = ['Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']

# Dictionary untuk melacak status proses
process_status = {}

# ======================================================================
# FUNGSI WEB SOCKET EVENT HANDLERS
# ======================================================================
@socketio.on('connect')
def handle_connect():
    print('Client connected:', request.sid)
    emit('connected', {'message': 'Connected to Flask WebSocket'})

@socketio.on('disconnect')
def handle_disconnect():
    print('Client disconnected:', request.sid)

@socketio.on('subscribe_to_progress')
def handle_subscribe(data):
    surat_id = data.get('surat_id')
    if surat_id:
        socketio.server.enter_room(request.sid, f'surat_{surat_id}')
        emit('subscribed', {'surat_id': surat_id}, room=request.sid)
        print(f'Client {request.sid} subscribed to surat_id: {surat_id}')

# ======================================================================
# FUNGSI UNTUK MENGIRIM STATUS MELALUI WEB SOCKET
# ======================================================================
def send_status(surat_id, status, message=None, progress=None, file_paths=None):
    """Mengirim status update ke client melalui WebSocket"""
    data = {
        'surat_id': surat_id,
        'status': status,
        'timestamp': datetime.now().isoformat()
    }
    
    if message:
        data['message'] = message
    if progress is not None:
        data['progress'] = progress
    if file_paths:
        data['file_paths'] = file_paths
    
    # Update status terbaru
    process_status[surat_id] = data
    
    # Kirim ke room yang sesuai
    socketio.emit('generation_status', data, room=f'surat_{surat_id}')
    print(f"[{surat_id}] Status update: {status} - {message}")

# ======================================================================
# FUNGSI BACKGROUND TASK DENGAN PROGRESS UPDATE
# ======================================================================
def background_task(surat_data):
    """Fungsi background dengan progress reporting melalui WebSocket"""
    surat_id = str(surat_data.get('id_surat_tugas_promotor', 'UNKNOWN'))
    laravel_upload_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/surat-promotor/upload-final'
    
    filename_docx = None
    filename_pdf = None
    
    try:
        # Update status: Memulai proses
        send_status(surat_id, 'processing', 'Memulai proses pembuatan surat...', 10)
        
        # 1. Persiapan data
        pythoncom.CoInitialize()
        
        def format_date(date_str):
            """Fungsi untuk memformat tanggal dari ISO format ke format Indonesia."""
            if not date_str: 
                return ""
            try:
                if isinstance(date_str, str):
                    if 'T' in date_str:
                        date_str = date_str.split('T')[0]
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                else:
                    date_obj = date_str
                return f"{date_obj.day} {bulan[date_obj.month - 1]} {date_obj.year}"
            except (ValueError, TypeError, IndexError) as e:
                print(f"Error formatting date: {date_str}, Error: {e}")
                return str(date_str)

        # Process penempatan
        penempatan_raw = surat_data.get('penempatan', '[]')
        penempatan_list = []
        try:
            parsed_data = json.loads(penempatan_raw) if isinstance(penempatan_raw, str) else penempatan_raw
            if isinstance(parsed_data, list):
                 penempatan_list = [str(item).strip() for item in parsed_data if str(item).strip()]
            else:
                 penempatan_list = [str(parsed_data)]
        except (json.JSONDecodeError, TypeError):
            penempatan_list = [item.strip() for item in str(penempatan_raw).split(',') if item.strip()]

        penempatan_text = "-"
        if penempatan_list:
            penempatan_text = penempatan_list[0] if len(penempatan_list) == 1 else "\n".join(f"{i}. {item}" for i, item in enumerate(penempatan_list, 1))

        nama_kandidat = surat_data.get('nama_kandidat', '-').strip() or '-'
        
        send_status(surat_id, 'processing', 'Memproses data template...', 30)
        
        # 2. Manipulasi dokumen Word
        file_template_path = os.getenv('TEMPLATE_PROMOTOR_PATH')
        if not os.path.exists(file_template_path):
            raise FileNotFoundError(f"Template tidak ditemukan di: {file_template_path}")
        
        document = Document(file_template_path)
        
        replacements = {
            '__PENEMPATAN__': penempatan_text,
            '__DATATGLSURATPEMBUATAN__': format_date(surat_data.get('tgl_surat_pembuatan')),
            '__NAMAKANDIDAT__': nama_kandidat,
            '__DATATGLPENUGASAN__': format_date(surat_data.get('tgl_penugasan')),
        }
        
        for key, value in replacements.items():
            for p in document.paragraphs: 
                p.text = p.text.replace(key, str(value))
            for t in document.tables:
                for r in t.rows:
                    for c in r.cells: 
                        c.text = c.text.replace(key, str(value))
        
        send_status(surat_id, 'processing', 'Menyimpan dokumen Word...', 50)
        
        # 3. Simpan file DOCX
        base_filename = f'Surat_Penempatan_Promotor_{nama_kandidat}_{surat_id}'
        safe_filename = re.sub(r'[^\w\-_\. ]', '_', base_filename)
        dirname = os.path.dirname(os.path.abspath(__file__))
        temp_dir = os.path.join(dirname, 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        filename_docx = os.path.join(temp_dir, f'{safe_filename}.docx')
        filename_pdf = os.path.join(temp_dir, f'{safe_filename}.pdf')
        
        document.save(filename_docx)
        
        send_status(surat_id, 'processing', 'Mengonversi ke PDF...', 70)
        
        # 4. Konversi ke PDF menggunakan docx2pdf
        try:
            convert(filename_docx, filename_pdf)
            send_status(surat_id, 'processing', 'Konversi PDF berhasil!', 80)
        except Exception as pdf_error:
            # Fallback: jika docx2pdf gagal, coba gunakan metode lain
            send_status(surat_id, 'processing', 'Menggunakan metode alternatif untuk konversi PDF...', 75)
            try:
                # Alternatif: menggunakan LibreOffice via command line (jika tersedia)
                import subprocess
                result = subprocess.run([
                    'soffice', '--headless', '--convert-to', 'pdf', 
                    '--outdir', temp_dir, filename_docx
                ], capture_output=True, text=True, timeout=60)
                
                if result.returncode != 0:
                    raise Exception(f"LibreOffice conversion failed: {result.stderr}")
                    
            except Exception as fallback_error:
                raise Exception(f"Konversi PDF gagal: {pdf_error}. Fallback juga gagal: {fallback_error}")
        
        send_status(surat_id, 'processing', 'Mengupload file ke server...', 90)
        
        # 5. Upload ke Laravel
        with open(filename_docx, 'rb') as f_docx, open(filename_pdf, 'rb') as f_pdf:
            files = {
                'file_docx': (os.path.basename(filename_docx), f_docx, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                'file_pdf': (os.path.basename(filename_pdf), f_pdf, 'application/pdf')
            }
            payload = {'surat_id': surat_id}
            
            res = requests.post(laravel_upload_url, files=files, data=payload, timeout=30)
            res.raise_for_status()
            
            response_json = res.json()
            if response_json.get('success'):
                send_status(surat_id, 'completed', 'File berhasil digenerate dan diupload!', 100)
            else:
                raise Exception(f"Laravel error: {response_json.get('message')}")

    except Exception as e:
        error_msg = f"Error dalam proses: {str(e)}"
        print(f"[{surat_id}] ERROR: {error_msg}")
        traceback.print_exc()
        send_status(surat_id, 'error', error_msg)
        
    finally:
        # 6. Cleanup
        try:
            if filename_docx and os.path.exists(filename_docx): 
                os.remove(filename_docx)
            if filename_pdf and os.path.exists(filename_pdf): 
                os.remove(filename_pdf)
            pythoncom.CoUninitialize()
        except Exception as cleanup_error:
            print(f"[{surat_id}] Cleanup error: {cleanup_error}")

# ======================================================================
# ENDPOINT UTAMA DENGAN WEB SOCKET
# ======================================================================
@app.route('/generate/surat/promotor', methods=['POST'])
def generate_surat_promotor():
    try:
        if not request.json or 'surat_data' not in request.json:
            return jsonify({'status': 'error', 'message': 'Payload JSON atau surat_data tidak ditemukan.'}), 400
        
        surat_data = request.json['surat_data']
        surat_id = str(surat_data.get('id_surat_tugas_promotor', 'UNKNOWN'))
        
        # Cek apakah proses sudah berjalan untuk surat_id ini
        if surat_id in process_status and process_status[surat_id].get('status') == 'processing':
            return jsonify({
                'status': 'processing', 
                'message': 'Proses untuk ID ini sudah berjalan.',
                'websocket_required': True
            })
        
        # Mulai background task tanpa parameter sid
        socketio.start_background_task(background_task, surat_data)
        
        return jsonify({
            'status': 'success', 
            'message': 'Proses pembuatan file dimulai.',
            'surat_id': surat_id,
            'websocket_required': True
        })
        
    except Exception as e:
        error_msg = f"Error memulai proses: {str(e)}"
        print(f"ERROR in generate_surat_promotor: {error_msg}")
        return jsonify({'status': 'error', 'message': error_msg}), 500

# ======================================================================
# ENDPOINT UNTUK MENGECEK STATUS (Fallback)
# ======================================================================
@app.route('/check/status/<surat_id>', methods=['GET'])
def check_status(surat_id):
    """Endpoint fallback untuk cek status (jika WebSocket tidak available)"""
    status = process_status.get(surat_id, {'status': 'not_found'})
    return jsonify(status)

# ======================================================================
# ENDPOINT UNTUK MENDAPATKAN HISTORY STATUS
# ======================================================================
@app.route('/status/history', methods=['GET'])
def status_history():
    """Mendapatkan history status semua proses"""
    return jsonify(process_status)

if __name__ == '__main__':
    print("Starting Flask server with WebSocket support...")
    socketio.run(app, debug=True, port=5000, host='0.0.0.0')