import pythoncom
import comtypes.client
import os
import sys
import asyncio
import threading
import time
import tempfile
import subprocess
import traceback
from docx import Document
from flask import Flask, jsonify, request
from flask_cors import CORS
from flask_socketio import SocketIO, emit
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import mysql.connector
from dotenv import load_dotenv
import re
from docx2pdf import convert
from docx.shared import Pt
import json

app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*")
load_dotenv()
wdFormatPDF = 17
bulan = ['Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']
processesId = []
process_status = {}

mydb = mysql.connector.connect(
  host=os.getenv('DB_HOST'),
  user=os.getenv('DB_USERNAME'),
  password=os.getenv('DB_PASSWORD'),
  database=os.getenv('DB_DATABASE')
)

def background_generate_file(filename_docx, filename_pdf, laravel_url, surat_id, table_name, id_column_name):
    asyncio.run(generateFile(filename_docx, filename_pdf, laravel_url, surat_id, table_name, id_column_name))

async def generateFile(filename_docx, filename_pdf, laravel_url, id_surat, table_name, id_column_name):
    f = open(filename_docx, 'rb')
    pdf = open(filename_pdf, 'rb')
    
    files = {
        'file_docx': (filename_docx, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
        'file_pdf': (filename_pdf, pdf, 'application/pdf')
    }

    res = requests.post(laravel_url, files=files, stream=True)
    dataJsonFromRes = res.json()

    mycursor = mydb.cursor()
    mycursor.execute(
    f"UPDATE {table_name} "f"SET file_path_docx = '{dataJsonFromRes['files']['docx']}', "f"file_path_pdf = '{dataJsonFromRes['files']['pdf']}' "f"WHERE {id_column_name} = '{id_surat}'")
    mydb.commit()
    f.close()
    pdf.close()
    t = threading.Thread(
        target=sendStatusProcess,
        args=(id_surat, False)
    )
    t.start()

    try:
        processesId.remove(id_surat)
    except:
        pass

    if os.path.exists(filename_docx):
        print(filename_docx)
        os.remove(filename_docx)

    if os.path.exists(filename_pdf):
        print(filename_pdf)
        os.remove(filename_pdf)

@socketio.on("connect")
def handle_connect():
    print("Client terhubung")
    emit("message", "Koneksi WebSocket Berhasil!", broadcast=True)

@socketio.on('get_info_process')
def getInfoProcess(data):
    print('Get Info Process....')
    laravel_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/get/info/file'
    requestPDF = requests.post(laravel_url, {
        'id': data,
        'type': 'pdf'
    })
    requestDOCX = requests.post(laravel_url, {
        'id': data,
        'type': 'docx'
    })

    if requestPDF.json()['status'] and requestDOCX.json()['status']:
        emit('fetch_status', {
            'status': True,
            'id': data,
        }, broadcast=True)
    else:
        emit('fetch_status', {
            'status': False,
            'id': data
        }, broadcast=True)

def sendStatusProcess(id, status):
    socketio.emit('send_status_process', {
        'status': status,
        'id': id
    })
    print('Send Status Process....', status)

@socketio.on('connect_after_fetch_table')
def connect_after(data):
    emit('connect_after', True)
    
@socketio.on('subscribe_to_progress')
def handle_subscribe_promotor(data):
    surat_id = data.get('surat_id')
    if surat_id:
        socketio.server.enter_room(request.sid, f'surat_{surat_id}')
        emit('subscribed', {'surat_id': surat_id}, room=request.sid)
        print(f'Client {request.sid} subscribed to surat_id: {surat_id}')

@app.route('/check/generate/run', methods=['POST'])
def checkGenerateFiles():
    for id in processesId:
        if id == str(request.json['id']):
            return {
                'status': True
            }
        
    return {
        'status': False
    }

@app.route('/generate/surat/penggati/driver', methods=['POST'])
def surat_penggati_driver():
    t = threading.Thread(
        target=sendStatusProcess,
        args=(request.json['id_surat_tugas'], True)
    )
    t.start()

    laravel_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/send/surat/pengganti/driver'
    file_template_path = 'Contoh Template/template_surat_pengganti_driver.docx'
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    now = datetime.now()
    document = Document(file_template_path)

    processesId.append(request.json['id_surat_tugas'])

    mycursor = mydb.cursor()
    sqlStr = f"SELECT * FROM surat_tugas_pengganti_driver WHERE id_surat_tugas = '{request.json['id_surat_tugas']}'"
    mycursor.execute(sqlStr)
    myresult = mycursor.fetchone()

    print(myresult)
    if myresult != None:

        """
        __TANGGALPEMBUATAN__
        __NAMAKANDIDAT__
        __NIKKANDIDAT__
        __JABATANKANDIDAT__
        __PENGGANTIKANDIDAT__
        __TANGGALMULAI__
        __TANGGALSELESAI__
        __DAERAHBANK__
        """

        for i in range(len(document.paragraphs)):
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__TANGGALPEMBUATAN__', f'{str(myresult[8]).split("-")[2]} {bulan[int(str(myresult[8]).split("-")[1]) - 1]} {str(myresult[8]).split("-")[0]}')
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__NAMAKANDIDAT__', myresult[1])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__NIKKANDIDAT__', myresult[2])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__JABATANKANDIDAT__', myresult[3])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__PENGGANTIKANDIDAT__', myresult[4])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__DAERAHBANK__', myresult[5])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__TANGGALMULAI__', f'{str(myresult[6]).split("-")[2]} {bulan[int(str(myresult[6]).split("-")[1]) - 1]} {str(myresult[6]).split("-")[0]}')
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__TANGGALSELESAI__', f'{str(myresult[7]).split("-")[2]} {bulan[int(str(myresult[7]).split("-")[1]) - 1]} {str(myresult[7]).split("-")[0]}')

        table = document.tables[0]
        cell1 = table.rows[1].cells[0]
        cell2 = table.rows[3].cells[0]
        cell3 = table.rows[4].cells[0]

        for para in cell1.paragraphs:
            for run in para.runs:
                run.bold = True

        for para in cell2.paragraphs:
            for run in para.runs:
                run.bold = True
                run.underline = True

        for para in cell3.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.italic = True
        

        document.save(f'Surat Tugas_{myresult[1]}.docx')
        dirname = os.path.dirname(__file__)
        filename_docx = os.path.join(dirname, f'Surat Tugas_{myresult[1]}.docx')
        filename_pdf = os.path.join(dirname, f'Surat Tugas_{myresult[1]}.pdf')
        word.Visible = False
        doc = word.Documents.Open(filename_docx, ReadOnly=True)
        doc.SaveAs(filename_pdf, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()

        # TODO
        t = threading.Thread(
            target=background_generate_file,
            args=(filename_docx, filename_pdf, laravel_url, request.json['id_surat_tugas'], request.json['table'], 'id_surat_tugas')
        )
        t.start()

        return jsonify({
            'status': 'success'
        })
    else:
        return jsonify({
            'status': 'error'
        })

@app.route('/generate/surat/tugas/mandiri', methods=['POST'])
def driver_mandiri():
    laravel_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/send/surat/pengganti/driver'
    file_template_path = 'Contoh Template/template_surat_penempatan_driver_mandiri.docx'
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    now = datetime.now()
    document = Document(file_template_path)

    mycursor = mydb.cursor()
    sqlStr = f"SELECT * FROM surat_tugas_mandiri WHERE id_surat_penempatan = '{request.json['id_surat_penempatan']}'"
    mycursor.execute(sqlStr)
    myresult = mycursor.fetchone()

    print(myresult)
    if myresult is not None:

        """
        __NOMORSURAT__
        __TANGGALPEMBUATAN__
        __NAMAKANDIDAT__
        __JABATANKANDIDAT__
        __TANGGALPENEMPATAN__
        """

        for i in range(len(document.paragraphs)):
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__NOMORSURAT__', myresult[1])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__TANGGALPEMBUATAN__',f'{str(myresult[2]).split("-")[2]} {bulan[int(str(myresult[2]).split("-")[1]) - 1]} {str(myresult[2]).split("-")[0]}')
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__NAMAKANDIDAT__', myresult[3])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__JABATANKANDIDAT__', myresult[4])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__TANGGALPENEMPATAN__',f'{str(myresult[5]).split("-")[2]} {bulan[int(str(myresult[5]).split("-")[1]) - 1]} {str(myresult[5]).split("-")[0]}')

        table = document.tables[0]
        cell1 = table.rows[1].cells[0]
        cell2 = table.rows[3].cells[0]
        cell3 = table.rows[4].cells[0]

        for para in cell1.paragraphs:
            for run in para.runs:
                run.bold = True

        for para in cell2.paragraphs:
            for run in para.runs:
                run.bold = True
                run.underline = True

        for para in cell3.paragraphs:
            for run in para.runs:
                run.italic = True

        document.save(f'Surat Penempatan_{myresult[3]}.docx')
        dirname = os.path.dirname(__file__)
        filename_docx = os.path.join(dirname, f'Surat Penempatan_{myresult[3]}.docx')
        filename_pdf = os.path.join(dirname, f'Surat Penempatan_{myresult[3]}.pdf')
        word.Visible = False
        doc = word.Documents.Open(filename_docx, ReadOnly=True)
        doc.SaveAs(filename_pdf, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()

        t = threading.Thread(
        target=background_generate_file,
        args=(
        filename_docx, filename_pdf, laravel_url, request.json['id_surat_penempatan'], request.json['table'], "id_surat_penempatan")
        )
        t.start()

        return jsonify({'status': 'success'})
    else:
        return jsonify({'status': 'error'})
    
# ======================================================================
# SURAT TUGAS PROMOTOR
# ======================================================================
def send_status_promotor(surat_id, status, message=None, progress=None, file_paths=None):
    """Mengirim status update ke client melalui WebSocket"""
    data = {
        'surat_id': surat_id,
        'status': status,
        'timestamp': datetime.now().isoformat()
    }
    
    if message:
        data['message'] = message
    if progress is not None:
        data['progress'] = progress
    if file_paths:
        data['file_paths'] = file_paths
    
    # Update status terbaru
    process_status[surat_id] = data
    
    # Kirim ke room yang sesuai
    socketio.emit('generation_status_promotor', data, room=f'surat_{surat_id}')
    print(f"[{surat_id}] Status update: {status} - {message}")

def background_task_promotor(surat_data):
    """Fungsi background dengan progress reporting melalui WebSocket"""
    surat_id = str(surat_data.get('id_surat_tugas_promotor', 'UNKNOWN'))
    laravel_upload_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/surat-promotor/upload-final'
    
    filename_docx = None
    filename_pdf = None
    
    try:
        # Update status: Memulai proses
        send_status_promotor(surat_id, 'processing', 'Memulai proses pembuatan surat...', 10)
        
        # 1. Persiapan data
        pythoncom.CoInitialize()
        
        def format_date_promotor(date_str):
            """Fungsi untuk memformat tanggal dari ISO format ke format Indonesia."""
            if not date_str: 
                return ""
            try:
                if isinstance(date_str, str):
                    if 'T' in date_str:
                        date_str = date_str.split('T')[0]
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                else:
                    date_obj = date_str
                return f"{date_obj.day} {bulan[date_obj.month - 1]} {date_obj.year}"
            except (ValueError, TypeError, IndexError) as e:
                print(f"Error formatting date: {date_str}, Error: {e}")
                return str(date_str)

        # Process penempatan
        penempatan_raw = surat_data.get('penempatan', '[]')
        penempatan_list = []
        try:
            parsed_data = json.loads(penempatan_raw) if isinstance(penempatan_raw, str) else penempatan_raw
            if isinstance(parsed_data, list):
                 penempatan_list = [str(item).strip() for item in parsed_data if str(item).strip()]
            else:
                 penempatan_list = [str(parsed_data)]
        except (json.JSONDecodeError, TypeError):
            penempatan_list = [item.strip() for item in str(penempatan_raw).split(',') if item.strip()]

        penempatan_text = "-"
        if penempatan_list:
            penempatan_text = penempatan_list[0] if len(penempatan_list) == 1 else "\n".join(f"{i}. {item}" for i, item in enumerate(penempatan_list, 1))

        nama_kandidat = surat_data.get('nama_kandidat', '-').strip() or '-'
        
        send_status_promotor(surat_id, 'processing', 'Memproses data template...', 30)
        
        # 2. Manipulasi dokumen Word
        file_template_path = os.getenv('TEMPLATE_PROMOTOR_PATH')
        if not os.path.exists(file_template_path):
            raise FileNotFoundError(f"Template tidak ditemukan di: {file_template_path}")
        
        document = Document(file_template_path)
        
        replacements = {
            '__PENEMPATAN__': penempatan_text,
            '__DATATGLSURATPEMBUATAN__': format_date_promotor(surat_data.get('tgl_surat_pembuatan')),
            '__NAMAKANDIDAT__': nama_kandidat,
            '__DATATGLPENUGASAN__': format_date_promotor(surat_data.get('tgl_penugasan')),
        }
        
        for key, value in replacements.items():
            for p in document.paragraphs: 
                p.text = p.text.replace(key, str(value))
            for t in document.tables:
                for r in t.rows:
                    for c in r.cells: 
                        c.text = c.text.replace(key, str(value))
        
        send_status_promotor(surat_id, 'processing', 'Menyimpan dokumen Word...', 50)
        
        # 3. Simpan file DOCX
        base_filename = f'Surat_Penempatan_Promotor_{nama_kandidat}_{surat_id}'
        safe_filename = re.sub(r'[^\w\-_\. ]', '_', base_filename)
        dirname = os.path.dirname(os.path.abspath(__file__))
        temp_dir = os.path.join(dirname, 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        filename_docx = os.path.join(temp_dir, f'{safe_filename}.docx')
        filename_pdf = os.path.join(temp_dir, f'{safe_filename}.pdf')
        
        document.save(filename_docx)
        
        send_status_promotor(surat_id, 'processing', 'Mengonversi ke PDF...', 70)
        
        # 4. Konversi ke PDF menggunakan docx2pdf
        try:
            convert(filename_docx, filename_pdf)
            send_status_promotor(surat_id, 'processing', 'Konversi PDF berhasil!', 80)
        except Exception as pdf_error:
            # Fallback: jika docx2pdf gagal, coba gunakan metode lain
            send_status_promotor(surat_id, 'processing', 'Menggunakan metode alternatif untuk konversi PDF...', 75)
            try:
                # Alternatif: menggunakan LibreOffice via command line (jika tersedia)
                result = subprocess.run([
                    'soffice', '--headless', '--convert-to', 'pdf', 
                    '--outdir', temp_dir, filename_docx
                ], capture_output=True, text=True, timeout=60)
                
                if result.returncode != 0:
                    raise Exception(f"LibreOffice conversion failed: {result.stderr}")
                    
            except Exception as fallback_error:
                raise Exception(f"Konversi PDF gagal: {pdf_error}. Fallback juga gagal: {fallback_error}")
        
        send_status_promotor(surat_id, 'processing', 'Mengupload file ke server...', 90)
        
        # 5. Upload ke Laravel
        with open(filename_docx, 'rb') as f_docx, open(filename_pdf, 'rb') as f_pdf:
            files = {
                'file_docx': (os.path.basename(filename_docx), f_docx, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                'file_pdf': (os.path.basename(filename_pdf), f_pdf, 'application/pdf')
            }
            payload = {'surat_id': surat_id}
            
            res = requests.post(laravel_upload_url, files=files, data=payload, timeout=30)
            res.raise_for_status()
            
            response_json = res.json()
            if response_json.get('success'):
                send_status_promotor(surat_id, 'completed', 'File berhasil digenerate dan diupload!', 100)
            else:
                raise Exception(f"Laravel error: {response_json.get('message')}")

    except Exception as e:
        error_msg = f"Error dalam proses: {str(e)}"
        print(f"[{surat_id}] ERROR: {error_msg}")
        traceback.print_exc()
        send_status_promotor(surat_id, 'error', error_msg)
        
    finally:
        # 6. Cleanup
        try:
            if filename_docx and os.path.exists(filename_docx): 
                os.remove(filename_docx)
            if filename_pdf and os.path.exists(filename_pdf): 
                os.remove(filename_pdf)
            pythoncom.CoUninitialize()
        except Exception as cleanup_error:
            print(f"[{surat_id}] Cleanup error: {cleanup_error}")

@app.route('/generate/surat/promotor', methods=['POST'])
def generate_surat_promotor():
    try:
        if not request.json or 'surat_data' not in request.json:
            return jsonify({'status': 'error', 'message': 'Payload JSON atau surat_data tidak ditemukan.'}), 400
        
        surat_data = request.json['surat_data']
        surat_id = str(surat_data.get('id_surat_tugas_promotor', 'UNKNOWN'))
        
        # Cek apakah proses sudah berjalan untuk surat_id ini
        if surat_id in process_status and process_status[surat_id].get('status') == 'processing':
            return jsonify({
                'status': 'processing', 
                'message': 'Proses untuk ID ini sudah berjalan.',
                'websocket_required': True
            })
        
        # Mulai background task tanpa parameter sid
        socketio.start_background_task(background_task_promotor, surat_data)
        
        return jsonify({
            'success': True,
            'status': 'success', 
            'message': 'Proses pembuatan file dimulai.',
            'surat_id': surat_id,
            'websocket_required': True
        })
        
    except Exception as e:
        error_msg = f"Error memulai proses: {str(e)}"
        print(f"ERROR in generate_surat_promotor: {error_msg}")
        return jsonify({
            'success': False,
            'status': 'error',
            'message': error_msg
        }), 500

@app.route('/check/status/promotor/<surat_id>', methods=['GET'])
def check_status_promotor(surat_id):
    """Endpoint fallback untuk cek status (jika WebSocket tidak available)"""
    status = process_status.get(surat_id, {'status': 'not_found'})
    return jsonify(status)

@app.route('/status/history/promotor', methods=['GET'])
def status_history_promotor():
    """Mendapatkan history status semua proses"""
    return jsonify(process_status)

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        'status': 'healthy',
        'service': 'flask-doc-generator',
        'timestamp': datetime.now().isoformat()
    })

if __name__ == '__main__':
    print("Starting Flask server with WebSocket support...")
    socketio.run(app, debug=True, host='0.0.0.0', port=5000)
