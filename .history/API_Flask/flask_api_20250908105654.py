import pythoncom
import comtypes.client
import os
import sys
import asyncio
import threading
import time
from docx import Document
from flask import Flask, jsonify, request
from flask_cors import CORS
from flask_socketio import SocketIO, emit
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import mysql.connector
from dotenv import load_dotenv
import re
from docx.shared import Pt
import json

app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*")
load_dotenv()
wdFormatPDF = 17
bulan = ['Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']
processesId = []

mydb = mysql.connector.connect(
  host=os.getenv('DB_HOST'),
  user=os.getenv('DB_USERNAME'),
  password=os.getenv('DB_PASSWORD'),
  database=os.getenv('DB_DATABASE')
)

def background_generate_file(filename_docx, filename_pdf, laravel_url, surat_id, table_name, id_column_name):
    asyncio.run(generateFile(filename_docx, filename_pdf, laravel_url, surat_id, table_name, id_column_name))

async def generateFile(filename_docx, filename_pdf, laravel_url, id_surat, table_name, id_column_name):
    f = open(filename_docx, 'rb')
    pdf = open(filename_pdf, 'rb')
    
    files = {
        'file_docx': (filename_docx, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
        'file_pdf': (filename_pdf, pdf, 'application/pdf')
    }

    res = requests.post(laravel_url, files=files, stream=True)
    print(res.text())
    
    dataJsonFromRes = res.json()

    mycursor = mydb.cursor()
    mycursor.execute(
    f"UPDATE {table_name} "f"SET file_path_docx = '{dataJsonFromRes['files']['docx']}', "f"file_path_pdf = '{dataJsonFromRes['files']['pdf']}' "f"WHERE {id_column_name} = '{id_surat}'")
    mydb.commit()
    f.close()
    pdf.close()
    t = threading.Thread(
        target=sendStatusProcess,
        args=(id_surat, False)
    )
    t.start()

    try:
        processesId.remove(id_surat)
    except:
        pass

    if os.path.exists(filename_docx):
        print(filename_docx)
        os.remove(filename_docx)

    if os.path.exists(filename_pdf):
        print(filename_pdf)
        os.remove(filename_pdf)

@socketio.on("connect")
def handle_connect():
    print("Client terhubung")
    emit("message", "Koneksi WebSocket Berhasil!", broadcast=True)

@socketio.on('get_info_process')
def getInfoProcess(data):
    print('Get Info Process....')
    laravel_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/get/info/file'
    requestPDF = requests.post(laravel_url, {
        'id': data,
        'type': 'pdf'
    })
    requestDOCX = requests.post(laravel_url, {
        'id': data,
        'type': 'docx'
    })

    if requestPDF.json()['status'] and requestDOCX.json()['status']:
        emit('fetch_status', {
            'status': True,
            'id': data,
        }, broadcast=True)
    else:
        emit('fetch_status', {
            'status': False,
            'id': data
        }, broadcast=True)

def sendStatusProcess(id, status):
    socketio.emit('send_status_process', {
        'status': status,
        'id': id
    })
    print('Send Status Process....', status)

@socketio.on('connect_after_fetch_table')
def connect_after(data):
    emit('connect_after', True)

@app.route('/check/generate/run', methods=['POST'])
def checkGenerateFiles():
    for id in processesId:
        if id == str(request.json['id']):
            return {
                'status': True
            }
        
    return {
        'status': False
    }

@app.route('/generate/surat/penggati/driver', methods=['POST'])
def surat_penggati_driver():
    t = threading.Thread(
        target=sendStatusProcess,
        args=(request.json['id_surat_tugas'], True)
    )
    t.start()

    laravel_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/send/surat/pengganti/driver'
    file_template_path = 'Contoh Template/template_surat_pengganti_driver.docx'
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    now = datetime.now()
    document = Document(file_template_path)

    processesId.append(request.json['id_surat_tugas'])

    mycursor = mydb.cursor()
    sqlStr = f"SELECT * FROM surat_tugas_pengganti_driver WHERE id_surat_tugas = '{request.json['id_surat_tugas']}'"
    mycursor.execute(sqlStr)
    myresult = mycursor.fetchone()

    print(myresult)
    if myresult != None:

        """
        __TANGGALPEMBUATAN__
        __NAMAKANDIDAT__
        __NIKKANDIDAT__
        __JABATANKANDIDAT__
        __PENGGANTIKANDIDAT__
        __TANGGALMULAI__
        __TANGGALSELESAI__
        __DAERAHBANK__
        """

        for i in range(len(document.paragraphs)):
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__TANGGALPEMBUATAN__', f'{str(myresult[8]).split("-")[2]} {bulan[int(str(myresult[8]).split("-")[1]) - 1]} {str(myresult[8]).split("-")[0]}')
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__NAMAKANDIDAT__', myresult[1])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__NIKKANDIDAT__', myresult[2])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__JABATANKANDIDAT__', myresult[3])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__PENGGANTIKANDIDAT__', myresult[4])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__DAERAHBANK__', myresult[5])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__TANGGALMULAI__', f'{str(myresult[6]).split("-")[2]} {bulan[int(str(myresult[6]).split("-")[1]) - 1]} {str(myresult[6]).split("-")[0]}')
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__TANGGALSELESAI__', f'{str(myresult[7]).split("-")[2]} {bulan[int(str(myresult[7]).split("-")[1]) - 1]} {str(myresult[7]).split("-")[0]}')

        table = document.tables[0]
        cell1 = table.rows[1].cells[0]
        cell2 = table.rows[3].cells[0]
        cell3 = table.rows[4].cells[0]

        for para in cell1.paragraphs:
            for run in para.runs:
                run.bold = True

        for para in cell2.paragraphs:
            for run in para.runs:
                run.bold = True
                run.underline = True

        for para in cell3.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.italic = True
        

        document.save(f'Surat Tugas_{myresult[1]}.docx')
        dirname = os.path.dirname(__file__)
        filename_docx = os.path.join(dirname, f'Surat Tugas_{myresult[1]}.docx')
        filename_pdf = os.path.join(dirname, f'Surat Tugas_{myresult[1]}.pdf')
        word.Visible = False
        doc = word.Documents.Open(filename_docx, ReadOnly=True)
        doc.SaveAs(filename_pdf, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()

        # TODO
        t = threading.Thread(
            target=background_generate_file,
            args=(filename_docx, filename_pdf, laravel_url, request.json['id_surat_tugas'], request.json['table'], 'id_surat_tugas')
        )
        t.start()

        return jsonify({
            'status': 'success'
        })
    else:
        return jsonify({
            'status': 'error'
        })

@app.route('/generate/surat/tugas/mandiri', methods=['POST'])
def driver_mandiri():
    laravel_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/send/surat/pengganti/driver'
    file_template_path = 'Contoh Template/template_surat_penempatan_driver_mandiri.docx'
    pythoncom.CoInitialize()
    word = comtypes.client.CreateObject('Word.Application')
    now = datetime.now()
    document = Document(file_template_path)

    mycursor = mydb.cursor()
    sqlStr = f"SELECT * FROM surat_tugas_mandiri WHERE id_surat_penempatan = '{request.json['id_surat_penempatan']}'"
    mycursor.execute(sqlStr)
    myresult = mycursor.fetchone()

    print(myresult)
    if myresult is not None:

        id_surat = request.json['id_surat_penempatan']
        processesId.append(id_surat)

        t_status = threading.Thread(
            target=sendStatusProcess,
            args=(id_surat, True)
        )

        t_status.start()

        """
        __NOMORSURAT__
        __TANGGALPEMBUATAN__
        __NAMAKANDIDAT__
        __JABATANKANDIDAT__
        __TANGGALPENEMPATAN__
        """

        for i in range(len(document.paragraphs)):
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__NOMORSURAT__', myresult[1])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__TANGGALPEMBUATAN__',f'{str(myresult[2]).split("-")[2]} {bulan[int(str(myresult[2]).split("-")[1]) - 1]} {str(myresult[2]).split("-")[0]}')
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__NAMAKANDIDAT__', myresult[3])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__JABATANKANDIDAT__', myresult[4])
            document.paragraphs[i].text = document.paragraphs[i].text.replace('__TANGGALPENEMPATAN__',f'{str(myresult[5]).split("-")[2]} {bulan[int(str(myresult[5]).split("-")[1]) - 1]} {str(myresult[5]).split("-")[0]}')

        table = document.tables[0]
        cell1 = table.rows[1].cells[0]
        cell2 = table.rows[3].cells[0]
        cell3 = table.rows[4].cells[0]

        for para in cell1.paragraphs:
            for run in para.runs:
                run.bold = True

        for para in cell2.paragraphs:
            for run in para.runs:
                run.bold = True
                run.underline = True

        for para in cell3.paragraphs:
            for run in para.runs:
                run.italic = True

        document.save(f'Surat Penempatan_{myresult[3]}.docx')
        dirname = os.path.dirname(__file__)
        filename_docx = os.path.join(dirname, f'Surat Penempatan_{myresult[3]}.docx')
        filename_pdf = os.path.join(dirname, f'Surat Penempatan_{myresult[3]}.pdf')
        word.Visible = False
        doc = word.Documents.Open(filename_docx, ReadOnly=True)
        doc.SaveAs(filename_pdf, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()

        t = threading.Thread(
        target=background_generate_file,
        args=(
        filename_docx, filename_pdf, laravel_url, request.json['id_surat_penempatan'], request.json['table'], "id_surat_penempatan")
        )
        t.start()

        return jsonify({'status': 'success'})
    else:
        return jsonify({'status': 'error'})



        
    
# Surat Tugas Promotor Indosat
@app.route('/generate/surat/promotor', methods=['POST'])
def generate_surat_promotor():
    laravel_url = f'{os.getenv("LARAVEL_ENDPOINT")}/api/send/surat/promotor'
    file_template_path = 'Contoh Template/275-Surat Tugas Promotor Indosat-ahmad.docx'
    
    word = None
    doc = None
   
    try:
        # Initialize Word application
        pythoncom.CoInitialize()
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
       
        # Get data from database
        mycursor = mydb.cursor(dictionary=True)
       
        # Pastikan id_surat_tugas_promotor ada dalam request
        if 'id_surat_tugas_promotor' not in request.json:
            return jsonify({
                'status': 'error',
                'message': 'id_surat_tugas_promotor is required in request'
            })
       
        id_surat = request.json['id_surat_tugas_promotor']
        
        # Commit any pending transactions to ensure we get latest data
        mydb.commit()
        
        sqlStr = f"SELECT * FROM surat_tugas_promotor WHERE id_surat_tugas_promotor = '{id_surat}'"
       
        print(f"Executing SQL: {sqlStr}")
        mycursor.execute(sqlStr)
        result = mycursor.fetchone()

        if not result:
            return jsonify({
                'status': 'error',
                'message': f'Data not found for id: {id_surat}'
            })

        # Load template document
        if not os.path.exists(file_template_path):
            return jsonify({
                'status': 'error',
                'message': f'Template file not found: {file_template_path}'
            })

        document = Document(file_template_path)

        # Format date function
        def format_date(date_str):
            if not date_str:
                return ""
            try:
                if isinstance(date_str, str):
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                else:
                    date_obj = date_str
                return f"{date_obj.day} {bulan[date_obj.month-1]} {date_obj.year}"
            except:
                return ""

        # Process penempatan data - handle various formats
        penempatan = result.get('penempatan', '[]')
        print(f"Raw penempatan data: {penempatan} (type: {type(penempatan)})")
        
        try:
            # Handle different formats of penempatan data
            if isinstance(penempatan, str):
                # Try to parse as JSON first
                if penempatan.startswith('[') and penempatan.endswith(']'):
                    penempatan = json.loads(penempatan)
                elif penempatan.startswith('"') and penempatan.endswith('"'):
                    # Remove quotes and split by comma
                    penempatan = penempatan.strip('"').split(',')
                else:
                    # Split by comma as fallback
                    penempatan = [item.strip() for item in penempatan.split(',') if item.strip()]
           
            # Ensure we have a list
            if not isinstance(penempatan, list):
                penempatan = [str(penempatan)]
                
            # Clean each item
            penempatan = [str(item).strip().strip('"\'[]') for item in penempatan if item and str(item).strip()]
            
        except Exception as e:
            print(f"Error processing penempatan: {str(e)}")
            penempatan = ["-"]

        print(f"Processed penempatan: {penempatan}")

        # Format penempatan
        penempatan_text = ""
        if penempatan and any(item.strip() for item in penempatan):
            if len(penempatan) == 1:
                penempatan_text = penempatan[0]
            else:
                numbered_list = []
                for i, item in enumerate(penempatan, 1):
                    if item.strip():  # Only add non-empty items
                        numbered_list.append(f"{i}. {item.strip()}")
                
                if numbered_list:
                    penempatan_text = "\n".join(numbered_list)
                else:
                    penempatan_text = "-"
        else:
            penempatan_text = "-"

        # Get other data
        nama_kandidat = result.get('nama_kandidat', '-') or '-'
        tgl_surat_pembuatan = result.get('tgl_surat_pembuatan', '') or ''
        tgl_penugasan = result.get('tgl_penugasan', '') or ''

        print(f"Data used for generation:")
        print(f"  Nama: {nama_kandidat}")
        print(f"  Penempatan: {penempatan_text}")
        print(f"  Tgl Pembuatan: {tgl_surat_pembuatan}")
        print(f"  Tgl Penugasan: {tgl_penugasan}")

        # Replace placeholders in paragraphs
        for paragraph in document.paragraphs:
            original_text = paragraph.text
            new_text = original_text
            
            new_text = new_text.replace('__PENEMPATAN__', penempatan_text)
            new_text = new_text.replace('__DATATGLSURATPEMBUATAN__', format_date(tgl_surat_pembuatan))
            new_text = new_text.replace('__NAMAKANDIDAT__', nama_kandidat)
            new_text = new_text.replace('__DATATGLPENUGASAN__', format_date(tgl_penugasan))
            
            if new_text != original_text:
                paragraph.text = new_text
                print(f"Replaced in paragraph: {original_text} -> {new_text}")

        # Replace placeholders in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    new_text = original_text
                    
                    new_text = new_text.replace('__PENEMPATAN__', penempatan_text)
                    new_text = new_text.replace('__DATATGLSURATPEMBUATAN__', format_date(tgl_surat_pembuatan))
                    new_text = new_text.replace('__NAMAKANDIDAT__', nama_kandidat)
                    new_text = new_text.replace('__DATATGLPENUGASAN__', format_date(tgl_penugasan))
                    
                    if new_text != original_text:
                        cell.text = new_text
                        print(f"Replaced in table cell: {original_text} -> {new_text}")

        # Save documents
        print('Saving Docx...')
        output_filename = f'Surat_Penempatan_Promotor_{nama_kandidat}_{id_surat}.docx'
        safe_filename = re.sub(r'[^\w\-_\. ]', '_', output_filename)
        document.save(safe_filename)
       
        dirname = os.path.dirname(__file__)
        filename_docx = os.path.join(dirname, safe_filename)
        filename_pdf = os.path.join(dirname, f'Surat_Penempatan_Promotor_{nama_kandidat}_{id_surat}.pdf')
       
        # Convert to PDF
        word.Visible = False
        doc = word.Documents.Open(filename_docx, ReadOnly=True)
        print('Saving PDF...')
        doc.SaveAs(filename_pdf, FileFormat=wdFormatPDF)
        doc.Close()
        doc = None
        word.Quit()
        word = None
        print('Conversion done!')

        # Start background thread for file upload
        print('Start threading in background...')
        t = threading.Thread(
            target=background_generate_file,
            args=(
                filename_docx,
                filename_pdf,
                laravel_url,
                id_surat,
                request.json.get('table', 'surat_tugas_promotor'),
                "id_surat_tugas_promotor"
            )
        )
        t.daemon = True
        t.start()
        
        # Add to processes list
        processesId.append(id_surat)
        
        print('File generation process started successfully')
        return jsonify({'status': 'success'})
       
    except Exception as e:
        print(f"Error in generate_surat_promotor: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Clean up
        try:
            if doc:
                doc.Close()
            if word:
                word.Quit()
        except:
            pass
            
        return jsonify({
            'status': 'error',
            'message': str(e)
        })
    finally:
        try:
            pythoncom.CoUninitialize()
        except:
            pass

if __name__ == '__main__':
    app.run(debug=True)